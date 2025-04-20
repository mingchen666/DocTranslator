from docx import Document
import os

def read_paragraph_text(document, max_chars=1000):
    texts = []
    current_text = ""
    for paragraph in document.paragraphs:
        if current_text and len(current_text) + len(paragraph.text) <= max_chars:
            current_text += '\n' + paragraph.text
        else:
            if current_text:
                texts.append(current_text)
            current_text = paragraph.text

    if current_text:
        texts.append(current_text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_texts = read_cell_text(cell, max_chars)
                for cell_text in cell_texts:
                    if current_text and len(current_text) + len(cell_text) <= max_chars:
                        current_text += '\n' + cell_text
                    else:
                        if current_text:
                            texts.append(current_text)
                        current_text = cell_text

    if current_text:
        texts.append(current_text)

    return texts


def read_rune_text(document, max_chars=1000):
    texts = []
    current_text = ""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if current_text and len(current_text) + len(run.text) <= max_chars:
                current_text += run.text
            else:
                if current_text:
                    texts.append(current_text)
                current_text = run.text

        if len(paragraph.hyperlinks) > 0:
            for hyperlink in paragraph.hyperlinks:
                for run in hyperlink.runs:
                    if current_text and len(current_text) + len(run.text) <= max_chars:
                        current_text += run.text
                    else:
                        if current_text:
                            texts.append(current_text)
                        current_text = run.text

        if current_text:
            texts.append(current_text)
            current_text = ""

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_texts = read_cell_text(cell, max_chars)
                for cell_text in cell_texts:
                    if current_text and len(current_text) + len(cell_text) <= max_chars:
                        current_text += '\n' + cell_text
                    else:
                        if current_text:
                            texts.append(current_text)
                        current_text = cell_text

    if current_text:
        texts.append(current_text)

    return texts


def read_cell_text(cell, max_chars=1000):
    texts = []
    current_text = ""
    for paragraph in cell.paragraphs:
        if current_text and len(current_text) + len(paragraph.text) <= max_chars:
            current_text += '\n' + paragraph.text
        else:
            if current_text:
                texts.append(current_text)
            current_text = paragraph.text

    if current_text:
        texts.append(current_text)

    return texts


def main():
    # 替换为你的 Word 文件路径
    file_path = r"F:\桌面文件\composes测试\images\“泉涌情深：润兵甘泉的故事”.docx"

    try:
        document = Document(file_path)
    except Exception as e:
        print(f"无法访问该文档: {e}")
        return

    # 提取文本块
    texts_paragraph = read_paragraph_text(document)
    texts_rune = read_rune_text(document)

    # 打印分段结果
    print("### 分段结果 (read_paragraph_text) ###")
    for i, text in enumerate(texts_paragraph):
        print(f"Segment {i+1}:\n{text}\n{'-'*80}")

    print("\n### 分段结果 (read_rune_text) ###")
    for i, text in enumerate(texts_rune):
        print(f"Segment {i+1}:\n{text}\n{'-'*80}")


if __name__ == "__main__":
    main()
